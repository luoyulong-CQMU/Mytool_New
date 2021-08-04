import subprocess
import uiautomation
import os
import time
import threading
import tkinter as tk
from tkinter import messagebox, filedialog
from docx.oxml.ns import qn
from docx import Document



class PrindPainIdCard(object):
    """

    自动填写住院证的类
    """
    painName ="Name"
    painAge = "Age"
    painInDay = "PainInDay"
    bedIndex = "bedIndex"
    painId = "painId"
    duringDay = "duringDay"
    painDiag = "painDiag"
    painGender = "painGender"

    def __init__(self,name=None):
        """

        :param name: 医生姓名
        """
        self.DoctorName = name

    def draw_mainwindow(self,outmessage=None):
        self.out_message = outmessage
        top_level = tk.Toplevel()
        top_level.wm_attributes('-topmost', 1)
        top_level.title('完善信息')
        # top_level.geometry('300x150')
        top_level.resizable(width=False, height=False)
        screenwidth = top_level.winfo_screenwidth()
        screenheight = top_level.winfo_screenheight()
        width = 300
        height = 150
        alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
        top_level.geometry(alignstr)
        doctor_name = tk.StringVar(top_level)
        input_lab = tk.Label(top_level, text='输入医师姓名：', font=('Arial', 10), width=13, height=2)
        input_lab.pack()
        input_lab.place(x=6,y=7)
        input_box = tk.Entry(top_level, textvariable=doctor_name)
        doctor_name.set(' ')
        input_box.pack()
        input_box.place(x=10,y=40)
        sure_btn = tk.Button(top_level, text="确定", font=('Arial', 10), width=8, height=1,
                  command=lambda: self.get_doctor_name(doctor_name, top_level))
        sure_btn.pack()
        sure_btn.place(x=210, y=35)
        top_level.mainloop()

    def get_doctor_name(self,doctor_name, toplevel):
        self.DoctorName = doctor_name.get().strip()
        toplevel.destroy()
        # 开启多线程
        new_t = threading.Thread(target=self.start)
        new_t.setDaemon(True)
        new_t.start()

    def start(self):
        self.get_pain_message()
        self.getMainDiag()
        self.loadDocument()
        self.fillGender()
        self.fillName()
        self.fillBexIndex()
        self.fillPainId()
        self.fillAge()
        self.fillDiag()
        self.fillInDay()
        self.fillDoctorName()
        self.saveDocument()

    def get_pain_message(self):
        window = uiautomation.WindowControl(Name="中联基层医疗卫生机构综合信息系统")
        clientPanel = window.PaneControl(AutomationId="clientPanel")
        panelControl1 = clientPanel.PaneControl(AutomationId="panelControl1")
        tab_1 = panelControl1.TabControl(AutomationId="xTabMain")
        doctorStaiton = tab_1.PaneControl(Name="住院医生站")
        allStationBox = doctorStaiton.PaneControl(AutomationId="SmartForm").PaneControl(Name="The XtraLayoutControl")
        allStationBoxChildAll = allStationBox.GetChildren()
        painInfoForJudge = allStationBoxChildAll[2].GetChildren()[0].AutomationId
        try:
            self.out_message.set(f"AutomationId:{painInfoForJudge}")
        except:
            pass
        # print("AutomationId:",painInfoForJudge,type(painInfoForJudge))
        trueId = "460924"
        try:
            painInfoBox = allStationBoxChildAll[2]
            painInfos = painInfoBox.GetChildren()[0].PaneControl(
                AutomationId="lcRecordViewHolder",searchDepth=1).GetChildren()
        except:
            try:
                self.out_message.set("提示：切换至医嘱界面可提高检索效率！")
            except:
                pass
            # print("提示：切换至医嘱界面可提高检索效率！")
            painInfoBox = allStationBoxChildAll[0]
            painInfos = painInfoBox.GetChildren()[0].PaneControl(
                AutomationId="lcRecordViewHolder", searchDepth=1).GetChildren()
        # 0 过敏史
        # 1 姓名
        # 2 性别
        # 3 年龄
        # 4 入院日期
        # 5 病况
        # 6 护理级别
        # 7 床号
        # 8 住院天数
        # 9 医保类型
        # 10 住院号
        # 11 诊断
        # 12 住院次数
        # 13 保险号
        self.painName = painInfos[1].GetChildren()[0]
        self.painInDay = painInfos[4].GetChildren()[0]
        self.bedIndex = painInfos[7].GetChildren()[0]
        self.duringDay = painInfos[8].GetChildren()[0]
        self.painId = painInfos[10].GetChildren()[0]
        self.painDiag = painInfos[11].GetChildren()[0]
        self.painAge = painInfos[3].GetChildren()[0]
        self.painGender = painInfos[2].GetChildren()[0]
        self.getMainDiag()
        # print(self.painGender.Name)
        # print(self.painMainDiag)

    def getMainDiag(self):
        self.painMainDiag = self.painDiag.Name.split("：")[1]
        try:
            self.painMainDiag = self.painMainDiag.split(",")[0]
        except:
            self.out_message.set("格式化诊断失败！")


    def loadDocument(self):
        self.path_ = "config/inHospitalCard.docx"
        self.document=Document(self.path_)
        self.document.styles['Normal'].font.name = u'宋体'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.table = self.document.tables[0]

    def fillGender(self):
        genderText = self.table.cell(0, 1).text.replace("painGender",self.painGender.Name)
        self.table.cell(0,1).text = genderText
        self.table.cell(2, 2).text = genderText

    def fillName(self):
        nameText = self.table.cell(0,0).text.replace("painName",self.painName.Name)
        self.table.cell(0,0).text = nameText
        self.table.cell(2,0).text = nameText


    def fillBexIndex(self):
        bexIndex = self.table.cell(0,6).text.replace("painIndex",self.bedIndex.Name)
        self.table.cell(0,6).text = bexIndex

    def fillPainId(self):
        painIdText = self.table.cell(0,7).text.replace("住院号:painID",self.painId.Name)
        self.table.cell(0,7).text = painIdText
        # print(painIdText)

    def fillAge(self):
        painAgeText = self.table.cell(2,6).text.replace("painAge",self.painAge.Name)
        self.table.cell(2,6).text = painAgeText

    def fillDiag(self):
        diagText = self.table.cell(6,0).text.replace("Diag",self.painMainDiag)
        self.table.cell(6,0).text=diagText

    def fillInDay(self):
        inDayText = self.table.cell(6,4).text.replace("inDay",self.painInDay.Name)
        self.table.cell(6,4).text = inDayText

    def fillDoctorName(self):
        doctorText = self.table.cell(10,0).text.replace("DoctorName",self.DoctorName)
        self.table.cell(10,0).text = doctorText

    def saveDocument(self):
        path_ = 'D:\\住院证\\'
        if not os.path.exists(path_):
            self.out_message.set('保存路径不存在！创建中. . .')
            # print(f"保存路径不存在！创建中. . .")
            os.mkdir(path_)
        self.save_path = f"{path_}{self.painName.Name}.{self.painInDay.Name}.docx"
        self.document.save(self.save_path)
        self.out_message.set(f"文件已成功保存至：{self.save_path}")
        time.sleep(1)
        os.startfile(self.save_path)

if __name__ == '__main__':
    obj = PrindPainIdCard()
    obj.draw_mainwindow()
